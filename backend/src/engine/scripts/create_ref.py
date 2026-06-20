from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_reference_docx(filepath):
    doc = Document()
    
    # Paper size and Margins
    section = doc.sections[0]
    section.page_width = Cm(21.0) # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.25)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    
    # Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set Tab Stops for Normal style
    # We want tabs at ~4.5cm, ~9.1cm, ~13.6cm for 4 columns
    # python-docx paragraph format can set tab stops:
    p_fmt = style.paragraph_format
    # Clear existing tabs
    if p_fmt.tab_stops is not None:
        p_fmt.tab_stops.clear_all()
    
    # p_fmt.tab_stops.add_tab_stop(Cm(4.5))
    # p_fmt.tab_stops.add_tab_stop(Cm(9.1))
    # p_fmt.tab_stops.add_tab_stop(Cm(13.6))

    # Add a paragraph to keep it valid
    doc.add_paragraph("Reference")
    
    # Add a table to force python-docx to generate the Table Grid style
    doc.add_table(rows=1, cols=1, style="Table Grid")
    
    doc.save(filepath)

import os

if __name__ == "__main__":
    storage_path = os.getenv("IMG_STORAGE_PATH", os.path.join(os.path.dirname(__file__), "..", "..", "..", "storage"))
    ref_path = os.path.join(storage_path, "reference.docx")
    create_reference_docx(ref_path)
