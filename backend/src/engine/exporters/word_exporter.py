import os
import re
from typing import List, Dict, Any
from .common import fix_soft_newlines, get_svg_native_width_inches

def _render_word_question_body(contest: dict, questions: List[dict], include_solution: bool = True) -> List[str]:
    total_mc = sum(1 for q in questions if q.get('question_type') == 'mc')
    total_tf = sum(1 for q in questions if q.get('question_type') == 'tf')
    total_sa = sum(1 for q in questions if q.get('question_type') == 'sa')
    total_oe = sum(1 for q in questions if q.get('question_type') == 'oe')

    lines: List[str] = []
    question_counter = 1
    printed_mc = False
    printed_tf = False
    printed_sa = False
    printed_oe = False
    
    for q in questions:
        q_type = q.get('question_type')
        
        effective_type = q_type
        if q_type == 'st':
            for c in questions:
                if c.get('parent_id') == q['id']:
                    effective_type = c.get('question_type')
                    break
                    
        if effective_type == 'mc' and not printed_mc:
            lines.append(f"\\textbf{{PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.}} Thí sinh trả lời từ câu 1 đến câu {total_mc}. Mỗi câu hỏi thí sinh chỉ chọn một phương án.\n\n")
            printed_mc = True
            question_counter = 1
        elif effective_type == 'tf' and not printed_tf:
            lines.append(f"\\textbf{{PHẦN II. Câu trắc nghiệm đúng sai.}} Thí sinh trả lời từ câu 1 đến câu {total_tf}. Trong mỗi ý a), b), c), d) ở mỗi câu hỏi, thí sinh chọn đúng hoặc sai.\n\n")
            printed_tf = True
            question_counter = 1
        elif effective_type == 'sa' and not printed_sa:
            lines.append(f"\\textbf{{PHẦN III. Câu trắc nghiệm trả lời ngắn.}} Thí sinh trả lời từ câu 1 đến câu {total_sa}.\n\n")
            printed_sa = True
            question_counter = 1
        elif effective_type == 'oe' and not printed_oe:
            lines.append(f"\\textbf{{PHẦN IV. Câu tự luận.}} Thí sinh trả lời từ câu 1 đến câu {total_oe}.\n\n")
            printed_oe = True
            question_counter = 1
            
        def fix_mismatched_environments(text: str) -> str:
            # Sửa lỗi phổ biến của các công cụ OCR: \begin{table} nhưng đóng bằng \end{center}
            return re.sub(r'(\\begin\{table\}(?:(?!\\begin\{center\}).)*?\\end\{tabular\}\s*)\\end\{center\}', r'\1\\end{table}', text, flags=re.DOTALL)

        content = fix_mismatched_environments(fix_soft_newlines(q.get('content', '') or ''))
        solution = fix_mismatched_environments(fix_soft_newlines(q.get('solution', '') or ''))
        options = q.get('options', [])
        
        def replace_img(match):
            import urllib.parse
            url = urllib.parse.unquote(match.group(1)).split('?')[0]
            if "/static/images/" in url:
                rel_path = url.split("/static/images/")[-1]
            else:
                rel_path = url.split("/")[-1]
            
            img_storage_path = os.getenv("IMG_STORAGE_PATH", "./storage")
            local_path = os.path.join(img_storage_path, rel_path)
            abs_path = os.path.abspath(local_path).replace('\\', '/')
            
            scale_factor = 0.4
            match_name = rel_path.split('/')[-1]
            for img in q.get('images', []):
                sp = img.get('storage_path', '')
                if sp and sp.replace('\\', '/').endswith(match_name):
                    sc = img.get('img_scale')
                    if sc is not None:
                        try:
                            scale_factor = float(sc)
                        except (ValueError, TypeError):
                            pass
                    break
            
            if abs_path.lower().endswith('.svg'):
                pdf_path = abs_path.replace('.svg', '.pdf')
                png_path = abs_path.replace('.svg', '.png')
                if os.path.exists(pdf_path) and not os.path.exists(png_path):
                    import subprocess
                    try:
                        subprocess.run(["pdftocairo", "-png", "-singlefile", "-r", "300", pdf_path, png_path.replace('.png', '')], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    except Exception as e:
                        print(f"Failed to convert PDF to PNG for Word export: {e}")
                
                if os.path.exists(png_path):
                    abs_path = png_path
                    
                native_in = get_svg_native_width_inches(abs_path.replace('.png', '.svg'))
                final_width = native_in * scale_factor * 1.5
                return f"\n\n[CENTER]\\includegraphics[width={final_width:.2f}in]{{{abs_path}}}\n\n"
            
            return f"\n\n[CENTER]\\includegraphics[width={scale_factor:.2f}\\textwidth]{{{abs_path}}}\n\n"
            
        content = re.sub(r'!\[.*?\]\((.*?)\)', replace_img, content)
        solution = re.sub(r'!\[.*?\]\((.*?)\)', replace_img, solution)
        
        # Check parent info if stimulus
        if q_type == 'st':
            child_count = sum(1 for c in questions if c.get('parent_id') == q['id'])
            if child_count > 0:
                start_c = question_counter
                end_c = question_counter + child_count - 1
                lines.append(f"\\textit{{Sử dụng dữ kiện sau để trả lời từ Câu {start_c} đến Câu {end_c}:}}\n\n{content}\n\n")
            else:
                lines.append(f"\\textit{{Sử dụng dữ kiện sau:}}\n\n{content}\n\n")
            continue
            
        lines.append(f"\\textbf{{Câu {question_counter}:}} {content}\n\n")
        
        if q_type == 'mc':
            # Decide layout: 4, 2, or 1 per line based on length
            opts_text = []
            has_image = False
            max_len = 0
            for idx, opt in enumerate(options):
                label = chr(65 + idx)
                opt_content = fix_soft_newlines(opt.get('content', '') or '')
                if '![' in opt_content: has_image = True
                
                # Estimate visible length by stripping math and latex tags
                s = re.sub(r'!\[.*?\]\(.*?\)', '', opt_content)
                s = re.sub(r'\\[a-zA-Z]+', '', s)
                s = re.sub(r'[\$\{\}\\_^]', '', s)
                s = re.sub(r'\s+', ' ', s).strip()
                max_len = max(max_len, len(s))
                
                opt_content = re.sub(r'!\[.*?\]\((.*?)\)', replace_img, opt_content).strip()
                opts_text.append(f"\\textbf{{{label}.}} {opt_content}")
                
            layout_type = str(q.get('layout_type', ''))
            
            if layout_type == '1':
                for opt_str in opts_text: lines.append(f"{opt_str}\n\n")
            elif layout_type == '2':
                lines.append(f"{opts_text[0]}[TAB2]{opts_text[1]}\n\n")
                if len(opts_text) > 2:
                    lines.append(f"{opts_text[2]}[TAB2]{opts_text[3]}\n\n")
            elif layout_type == '4':
                lines.append("[TAB4]".join(opts_text) + "\n\n")
            else:
                if has_image or max_len > 35:
                    # 1 per line
                    for opt_str in opts_text:
                        lines.append(f"{opt_str}\n\n")
                elif max_len > 12:
                    # 2 per line
                    lines.append(f"{opts_text[0]}[TAB2]{opts_text[1]}\n\n")
                    if len(opts_text) > 2:
                        lines.append(f"{opts_text[2]}[TAB2]{opts_text[3]}\n\n")
                else:
                    # 4 per line
                    lines.append("[TAB4]".join(opts_text) + "\n\n")
            
            if include_solution:
                lines.append("[CENTER]\\textbf{Lời giải}\n\n")
                correct_label = "##"
                for idx, opt in enumerate(options):
                    if opt.get('is_correct'):
                        correct_label = chr(65 + idx)
                        break
                        
                if not solution.strip():
                    lines.append(f"\\textbf{{Chọn {correct_label}}}\n\n")
                else:
                    lines.append(f"{solution}\n\n\\textbf{{Chọn {correct_label}}}\n\n")
        elif q_type == 'tf':
            for idx, opt in enumerate(options):
                label = chr(97 + idx) # a, b, c, d
                opt_content = fix_soft_newlines(opt.get('content', '') or '')
                opt_content = re.sub(r'!\[.*?\]\((.*?)\)', replace_img, opt_content)
                lines.append(f"\\textbf{{{label})}} {opt_content}\n\n")
            
            if include_solution:
                lines.append("[CENTER]\\textbf{Lời giải}\n\n")
                if solution.strip():
                    lines.append(f"{solution}\n\n")
                for idx, opt in enumerate(options):
                    label = chr(97 + idx)
                    tf_status = "Đúng" if opt.get('is_correct') else "Sai"
                    explain = fix_soft_newlines(opt.get('explaination', '') or '')
                    explain = re.sub(r'!\[.*?\]\((.*?)\)', replace_img, explain)
                    lines.append(f"\\textbf{{{label}) {tf_status}.}} {explain}\n\n")
                
        elif q_type == 'sa':
            if include_solution:
                lines.append("[CENTER]\\textbf{Lời giải}\n\n")
                correct_ans = options[0].get('content', '') if options else ''
                lines.append(f"\\textbf{{Trả lời ngắn:}} {correct_ans}\n\n")
                if solution.strip():
                    lines.append(f"{solution}\n\n")
                
        elif q_type == 'oe':
            if include_solution:
                lines.append("[CENTER]\\textbf{Lời giải}\n\n")
                if solution.strip():
                    lines.append(f"{solution}\n\n")
                
        question_counter += 1
        
    lines.append("\n[CENTER]\\textbf{" + r"-\/" * 21 + " HẾT " + r"-\/" * 21 + "}\n\n")
    return lines


def _fmt_points(x: float) -> str:
    """Định dạng điểm kiểu Việt Nam: 14.0 -> '14,0', 0.25 -> '0,25'."""
    s = f"{x:.2f}"
    if s.endswith('0'):
        s = s[:-1]
    return s.replace('.', ',')


def _get_scoring_config(contest: dict) -> Dict[str, float]:
    import json
    sc = contest.get('scoring_config') if isinstance(contest, dict) else None
    if isinstance(sc, str):
        try:
            sc = json.loads(sc)
        except Exception:
            sc = {}
    if not isinstance(sc, dict):
        sc = {}
    return {
        'mc': float(sc.get('mc', 0.25) or 0.25),
        'tf': float(sc.get('tf', 1.0) or 1.0),
        'sa': float(sc.get('sa', 0.5) or 0.5),
        'oe': float(sc.get('oe', 1.0) or 1.0),
    }


def _render_answer_key_body(contest: dict, questions: List[dict], exam_title: str = "", department: str = "", subject: str = "", code: str = "000") -> List[str]:
    """Dựng bảng đáp án (hướng dẫn chấm) cho đề gốc theo mẫu THPT."""
    mc_list = [q for q in questions if q.get('question_type') == 'mc']
    tf_list = [q for q in questions if q.get('question_type') == 'tf']
    sa_list = [q for q in questions if q.get('question_type') == 'sa']
    oe_list = [q for q in questions if q.get('question_type') == 'oe']

    w = _get_scoring_config(contest)
    has_oe = len(oe_list) > 0

    lines: List[str] = []

    # Marker: post-process sẽ chèn bảng tiêu đề (giống phần đề) và đẩy sang trang mới
    lines.append("[ANSWERHEADER]\n\n")

    def mc_letter(q: dict) -> str:
        for idx, opt in enumerate(q.get('options', [])):
            if opt.get('is_correct'):
                return chr(65 + idx)
        return ''

    # ── Phần trắc nghiệm khách quan ──
    if has_oe:
        points_A = len(mc_list) * w['mc'] + len(tf_list) * w['tf'] + len(sa_list) * w['sa']
        lines.append(f"\\textbf{{PHẦN TRẮC NGHIỆM: {_fmt_points(points_A)} điểm}}\n\n")

    # I. Trắc nghiệm nhiều phương án lựa chọn — bảng số câu / đáp án, 12 câu mỗi hàng
    if mc_list:
        head = "I." if has_oe else "PHẦN I."
        lines.append(f"\\textbf{{{head} Câu trắc nghiệm nhiều phương án lựa chọn:}} Mỗi câu trả lời đúng thí sinh được {_fmt_points(w['mc'])} điểm.\n\n")
        items = [(i + 1, mc_letter(q)) for i, q in enumerate(mc_list)]
        per = 10
        ncol = min(per, len(items))
        lines.append('\\begin{tabular}{' + ('|c' * ncol) + '|}\n\\hline\n')
        for s in range(0, len(items), per):
            chunk = items[s:s + per]
            nums = [str(n) for n, _ in chunk] + [''] * (ncol - len(chunk))
            ans = [(a if a else '') for _, a in chunk] + [''] * (ncol - len(chunk))
            lines.append(' & '.join(nums) + ' \\\\\n\\hline\n')
            lines.append(' & '.join(ans) + ' \\\\\n\\hline\n')
        lines.append('\\end{tabular}\n\n')

    # II. Trắc nghiệm đúng sai — bảng dọc a/b/c/d, hai câu cạnh nhau
    if tf_list:
        head = "II." if has_oe else "PHẦN II."
        lines.append(f"\\textbf{{{head} Câu trắc nghiệm đúng sai.}}\n\n")
        lines.append(f"- Thí sinh chỉ lựa chọn chính xác 01 ý trong 1 câu hỏi được {_fmt_points(0.1 * w['tf'])} điểm.\n\n")
        lines.append(f"- Thí sinh chỉ lựa chọn chính xác 02 ý trong 1 câu hỏi được {_fmt_points(0.25 * w['tf'])} điểm.\n\n")
        lines.append(f"- Thí sinh chỉ lựa chọn chính xác 03 ý trong 1 câu hỏi được {_fmt_points(0.5 * w['tf'])} điểm.\n\n")
        lines.append(f"- Thí sinh lựa chọn chính xác cả 04 ý trong 1 câu hỏi được {_fmt_points(1.0 * w['tf'])} điểm.\n\n")

        # Bảng Đ/S được dựng bằng python-docx ở post-process (gộp ô cột "Câu" cho đẹp)
        lines.append("[TFTABLE]\n\n")

    # III. Trắc nghiệm trả lời ngắn — bảng số câu / đáp án, 8 câu mỗi hàng
    if sa_list:
        head = "III." if has_oe else "PHẦN III."
        lines.append(f"\\textbf{{{head} Câu trắc nghiệm trả lời ngắn:}} Mỗi câu trả lời đúng thí sinh được {_fmt_points(w['sa'])} điểm.\n\n")
        items = []
        for i, q in enumerate(sa_list):
            ans = q.get('options', [{}])[0].get('content', '') if q.get('options') else ''
            ans = str(ans).replace('\n', ' ').strip()
            items.append((i + 1, ans))
        per = 10
        ncol = min(per, len(items))
        lines.append('\\begin{tabular}{|c|' + ('c|' * ncol) + '}\n\\hline\n')
        for s in range(0, len(items), per):
            chunk = items[s:s + per]
            nums = [str(n) for n, _ in chunk] + [''] * (ncol - len(chunk))
            ans = [(a if a else '') for _, a in chunk] + [''] * (ncol - len(chunk))
            lines.append('Câu & ' + ' & '.join(nums) + ' \\\\\n\\hline\n')
            lines.append('Đáp án & ' + ' & '.join(ans) + ' \\\\\n\\hline\n')
        lines.append('\\end{tabular}\n\n')

    # ── Phần tự luận (ghi tổng điểm; lời giải nằm ở mục Lời giải chi tiết bên dưới) ──
    if has_oe:
        points_B = len(oe_list) * w['oe']
        lines.append(f"\\textbf{{PHẦN TỰ LUẬN: {_fmt_points(points_B)} điểm}}\n\n")

    # ── Lời giải chi tiết toàn bộ câu hỏi (nằm bên dưới bảng đáp án) ──
    lines.append("~\n\n")  # 1 dòng trống ngăn cách với bảng đáp án phía trên
    lines.append("[CENTER]\\textbf{LỜI GIẢI CHI TIẾT}\n\n")
    lines += _render_word_question_body(contest, questions, include_solution=True)

    return lines


def get_word_simplified_latex(contest: dict, questions: List[dict], exam_title: str = "", department: str = "", exam_type: str = "", subject: str = "", duration: int = 50, general_info: str = "", code: str = "000", include_solution: bool = True, dual_section: bool = False) -> str:
    lines = ["\\begin{document}\n\n"]
    if general_info:
        for line in general_info.strip().split('\n'):
            line = line.strip()
            if line:
                lines.append(f"{line}\n\n")

    if dual_section:
        # Phần 1: chỉ đề thi (không lời giải)
        lines += _render_word_question_body(contest, questions, include_solution=False)
        # Bảng đáp án (hướng dẫn chấm) — sẽ được đẩy sang trang mới ở post-process
        # qua marker [PAGEBREAK] (pandoc không dịch \newpage thành ngắt trang Word).
        lines += _render_answer_key_body(contest, questions, exam_title, department, subject, code)
    else:
        lines += _render_word_question_body(contest, questions, include_solution=include_solution)

    lines.append("\\end{document}")
    raw = "".join(lines)
    from .common import balance_latex_braces
    return balance_latex_braces(raw)

def post_process_word_docx(tmp_docx: str, original_code: str, department: str, exam_type: str, exam_title: str, contest: dict, subject: str, duration: int, general_info: str, dual_section: bool = False, questions: list = None):
    import os
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    
    def create_field(field_text, italic=False):
        def get_rPr():
            rPr = OxmlElement('w:rPr')
            if italic:
                rPr.append(OxmlElement('w:i'))
            return rPr
            
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r1 = OxmlElement('w:r')
        r1.append(get_rPr())
        r1.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' {field_text} '
        r2 = OxmlElement('w:r')
        r2.append(get_rPr())
        r2.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        r3 = OxmlElement('w:r')
        r3.append(get_rPr())
        r3.append(fldChar2)
        
        t = OxmlElement('w:t')
        t.text = '...'
        r4 = OxmlElement('w:r')
        r4.append(get_rPr())
        r4.append(t)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        r5 = OxmlElement('w:r')
        r5.append(get_rPr())
        r5.append(fldChar3)
        
        return [r1, r2, r3, r4, r5]
    
    if not os.path.exists(tmp_docx):
        return
        
    doc = docx.Document(tmp_docx)
    
    if not doc.paragraphs:
        return
    first_p = doc.paragraphs[0]
    
    def remove_borders(table):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            tblBorders = tblPr[0].xpath('w:tblBorders')
            if tblBorders:
                tblPr[0].remove(tblBorders[0])
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), 'none')
            borders.append(e)
        if tblPr:
            tblPr[0].append(borders)
    
    header_tables = []

    def build_header(anchor_p, code, is_answer=False):
        # Top table: 1 row, 2 cols
        t1 = doc.add_table(rows=1, cols=2)
        remove_borders(t1)
        t1.autofit = False
        t1.columns[0].width = Cm(6.25)
        t1.columns[1].width = Cm(12.5)

        # Left cell
        c = t1.cell(0, 0)
        c.width = Cm(6.25)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(department or 'BỘ GIÁO DỤC VÀ ĐÀO TẠO').bold = True
        p.add_run('\n')
        p.add_run(exam_type or 'ĐỀ THI CHÍNH THỨC')
        p.add_run('\n')
        r1 = p.add_run('(Đáp án có ' if is_answer else '(Đề thi có ')
        r1.italic = True
        for el in create_field(r'NUMPAGES \# "00"', italic=True):
            p._p.append(el)
        r2 = p.add_run(' trang)')
        r2.italic = True
        p.paragraph_format.space_after = Pt(1)

        # Right cell
        c = t1.cell(0, 1)
        c.width = Cm(12.5)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = exam_title if exam_title else contest.get('title', 'Đề thi')
        p.add_run(title).bold = True
        p.add_run('\n')
        p.add_run('Môn thi: ' + (subject or '...')).bold = True
        p.add_run('\n')
        r_dur = p.add_run('Thời gian làm bài: ' + str(duration or 50) + ' phút, không kể thời gian phát đề')
        r_dur.italic = True
        p.paragraph_format.space_after = Pt(1)

        anchor_p._p.addprevious(t1._tbl)

        # Spacer giữa hai bảng tiêu đề
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(3)
        anchor_p._p.addprevious(spacer._p)

        # Bottom table: 1 row, 2 cols
        t2 = doc.add_table(rows=1, cols=2)
        remove_borders(t2)
        t2.autofit = False
        t2.columns[0].width = Cm(12.5)
        t2.columns[1].width = Cm(6.25)

        # Left cell — Họ tên / SBD (bỏ ở phần đáp án)
        c = t2.cell(0, 0)
        c.width = Cm(12.5)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if not is_answer:
            p.add_run('Họ, tên thí sinh: ........................................................................\n').bold = True
            p.add_run('Số báo danh: .............................................................................').bold = True

        # Right cell (Mã đề box)
        c = t2.cell(0, 1)
        c.width = Cm(6.25)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)

        vml_xml = f"""
        <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:o="urn:schemas-microsoft-com:office:office">
          <w:pict>
            <v:rect id="_x0000_s1026" style="width:130pt;height:18pt;v-text-anchor:middle;" fillcolor="white" strokecolor="black" strokeweight="0.75pt">
              <v:textbox inset="0,0,0,0">
                <w:txbxContent>
                  <w:p>
                    <w:pPr>
                      <w:jc w:val="center"/>
                      <w:spacing w:after="0"/>
                    </w:pPr>
                    <w:r>
                      <w:rPr><w:b/></w:rPr>
                      <w:t>Mã đề: {code}</w:t>
                    </w:r>
                  </w:p>
                </w:txbxContent>
              </v:textbox>
            </v:rect>
          </w:pict>
        </w:r>
        """
        r_shape = parse_xml(vml_xml)
        p._p.append(r_shape)

        anchor_p._p.addprevious(t2._tbl)
        header_tables.extend([t1, t2])

    # Tiêu đề phần đề thi
    build_header(first_p, original_code, is_answer=False)

    # Tiêu đề phần đáp án (bảng giống phần đề), đẩy sang trang mới
    for p_ah in list(doc.paragraphs):
        if '[ANSWERHEADER]' in p_ah.text:
            for r in p_ah.runs:
                r.text = r.text.replace('[ANSWERHEADER]', '')
            # Ngắt trang bằng cách thêm page-break vào cuối đoạn liền trước (dòng
            # HẾT của đề), tránh tạo một đoạn rỗng thừa ở đầu trang đáp án.
            prev = p_ah._p.getprevious()
            if prev is not None and prev.tag == qn('w:p'):
                r_br = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                r_br.append(br)
                prev.append(r_br)
            else:
                p_ah.paragraph_format.page_break_before = True
            build_header(p_ah, original_code, is_answer=True)
            break

    # Bảng đáp án Đúng/Sai — dựng bằng python-docx để gộp ô cột "Câu" cho đẹp
    def _fill_cell(cell, text, bold=False, center=False, vcenter=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        if vcenter:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def build_tf_table(anchor_p, tf_list):
        n = len(tf_list)
        if n == 0:
            return
        half = (n + 1) // 2
        table = doc.add_table(rows=1 + half * 4, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = ['a)', 'b)', 'c)', 'd)']

        def ds(q):
            return ['Đúng' if o.get('is_correct') else 'Sai' for o in q.get('options', [])]

        for j, h in enumerate(['Câu', 'Lệnh hỏi', 'Đáp án', 'Câu', 'Lệnh hỏi', 'Đáp án']):
            _fill_cell(table.cell(0, j), h, center=True)

        for i in range(half):
            base = 1 + i * 4
            left = tf_list[i]; ld = ds(left)
            ri = i + half
            right = tf_list[ri] if ri < n else None
            rd = ds(right) if right else []
            for r in range(4):
                _fill_cell(table.cell(base + r, 1), labels[r] if r < len(ld) else '', center=True)
                _fill_cell(table.cell(base + r, 2), ld[r] if r < len(ld) else '', center=True)
                _fill_cell(table.cell(base + r, 4), labels[r] if (right and r < len(rd)) else '', center=True)
                _fill_cell(table.cell(base + r, 5), rd[r] if (right and r < len(rd)) else '', center=True)
            # Gộp ô cột "Câu" (cột 0) cho 4 hàng a/b/c/d
            cl = table.cell(base, 0)
            for r in range(1, 4):
                cl = cl.merge(table.cell(base + r, 0))
            _fill_cell(cl, str(i + 1), center=True, vcenter=True)
            # Gộp ô cột "Câu" bên phải (cột 3)
            cr = table.cell(base, 3)
            for r in range(1, 4):
                cr = cr.merge(table.cell(base + r, 3))
            if right:
                _fill_cell(cr, str(ri + 1), center=True, vcenter=True)

        anchor_p._p.addprevious(table._tbl)

    if questions:
        tf_list = [q for q in questions if q.get('question_type') == 'tf']
        for p_tf in list(doc.paragraphs):
            if '[TFTABLE]' in p_tf.text:
                for r in p_tf.runs:
                    r.text = r.text.replace('[TFTABLE]', '')
                build_tf_table(p_tf, tf_list)
                break
    
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run('Trang ')
    for el in create_field('PAGE'): fp._p.append(el)
    fp.add_run('/')
    for el in create_field('NUMPAGES'): fp._p.append(el)
    fp.add_run(f' - Mã đề thi {original_code}')
    
    # Align all display math (oMathPara) to the left safely using iter
    for oMathPara in doc.element.iter(qn('m:oMathPara')):
        jc_found = False
        for jc in oMathPara.iter(qn('m:jc')):
            jc.set(qn('m:val'), 'left')
            jc_found = True
        
        if not jc_found:
            pr = None
            for child in oMathPara.iter(qn('m:oMathParaPr')):
                pr = child
                break
            if pr is None:
                pr = OxmlElement('m:oMathParaPr')
                oMathPara.insert(0, pr)
            jc_el = OxmlElement('m:jc')
            jc_el.set(qn('m:val'), 'left')
            pr.append(jc_el)
            
    # Set space_after = 0 and ensure math is left-aligned for ALL paragraphs globally
    for p_all in doc.paragraphs:
        p_all.paragraph_format.space_after = Pt(0)
        p_all.paragraph_format.first_line_indent = Pt(0)
        
        # Apply justify to regular text (except explicitly centered/left things)
        if p_all.alignment is None or p_all.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            p_all.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Clear default tab stops added by Pandoc
        p_all.paragraph_format.tab_stops.clear_all()
        
        # Add space before section headers
        if p_all.text.startswith('PHẦN I.') or p_all.text.startswith('PHẦN II.') or p_all.text.startswith('PHẦN III.') or p_all.text.startswith('PHẦN IV.'):
            p_all.paragraph_format.space_before = Pt(8)
            
        if p_all.text.startswith('Sử dụng dữ kiện sau'):
            p_all.paragraph_format.space_before = Pt(8)
            
        xml_str = p_all._p.xml
        if 'm:oMath' in xml_str or 'm:oMathPara' in xml_str:
            if p_all.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                p_all.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
        # Process markers
        has_center = False
        has_tab2 = False
        has_tab4 = False
        for r in p_all.runs:
            if '[CENTER]' in r.text:
                r.text = r.text.replace('[CENTER]', '')
                has_center = True
            if '[TAB2]' in r.text:
                r.text = r.text.replace('[TAB2]', '\t')
                has_tab2 = True
            if '[TAB4]' in r.text:
                r.text = r.text.replace('[TAB4]', '\t')
                has_tab4 = True

        if has_center:
            p_all.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if has_tab2:
            p_all.paragraph_format.tab_stops.add_tab_stop(Cm(9.125), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
        elif has_tab4:
            p_all.paragraph_format.tab_stops.add_tab_stop(Cm(4.5625), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p_all.paragraph_format.tab_stops.add_tab_stop(Cm(9.125), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            p_all.paragraph_format.tab_stops.add_tab_stop(Cm(13.6875), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            
    # Ensure tables have borders and are centered
    header_tbl_elems = [h._tbl for h in header_tables]
    for t_all in doc.tables:
        if t_all._tbl not in header_tbl_elems:
            t_all.style = 'Table Grid'
            t_all.autofit = True
            t_all.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Fit to page: chiều rộng bảng = 100% chiều rộng trang
            tblPr = t_all._element.tblPr
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')
        
        for row in t_all.rows:
            for cell in row.cells:
                for p_all in cell.paragraphs:
                    p_all.paragraph_format.space_after = Pt(0)
                    xml_str = p_all._p.xml
                    if 'm:oMath' in xml_str or 'm:oMathPara' in xml_str:
                        if p_all.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            p_all.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
    # Set document-level Math default justification to Left
    settings = doc.settings.element
    math_pr = settings.find(qn('m:mathPr'))
    if math_pr is None:
        math_pr = OxmlElement('m:mathPr')
        settings.append(math_pr)
    def_jc = math_pr.find(qn('m:defJc'))
    if def_jc is None:
        def_jc = OxmlElement('m:defJc')
        math_pr.append(def_jc)
    def_jc.set(qn('m:val'), 'left')
                    
    doc.save(tmp_docx)


import subprocess
import tempfile
import zipfile

def export_word(contest: dict, questions: list, code: str, exam_title: str, department: str, exam_type: str, subject: str, duration: int, general_info: str, include_solution: bool, zf: zipfile.ZipFile, dual_section: bool = False):
    word_tex = get_word_simplified_latex(contest, questions, exam_title, department, exam_type, subject, duration, general_info, code, include_solution, dual_section)
    with tempfile.NamedTemporaryFile(suffix=".tex", delete=False, mode='w', encoding='utf-8') as f_tmp:
        f_tmp.write(word_tex)
        tmp_tex = f_tmp.name
    tmp_docx = tmp_tex.replace(".tex", ".docx")
    try:
        cmd = ["pandoc", tmp_tex, "-f", "latex", "-o", tmp_docx, "--mathml"]
        
        sp = os.getenv("IMG_STORAGE_PATH", "./storage")
        ref_path = os.path.join(sp, "reference.docx")
        if not os.path.exists(ref_path):
            try:
                import importlib.util
                import sys
                script_path = os.path.join(os.path.dirname(__file__), "..", "scripts", "create_ref.py")
                spec = importlib.util.spec_from_file_location("create_ref", script_path)
                create_ref_module = importlib.util.module_from_spec(spec)
                sys.modules["create_ref"] = create_ref_module
                spec.loader.exec_module(create_ref_module)
                
                os.makedirs(sp, exist_ok=True)
                create_ref_module.create_reference_docx(ref_path)
            except Exception as e:
                print(f"Failed to create reference.docx: {e}")
                
        if os.path.exists(ref_path):
            cmd.extend(["--reference-doc", ref_path])
        subprocess.run(cmd, check=True, timeout=90, capture_output=True, text=True)
        if os.path.exists(tmp_docx):
            post_process_word_docx(tmp_docx, code, department, exam_type, exam_title, contest, subject, duration, general_info, dual_section=dual_section, questions=questions)
            zf.write(tmp_docx, f"{code}.docx")
    except subprocess.CalledProcessError as e:
        print(f"Pandoc error output: {e.stderr}")
    except Exception as e:
        print(f"Pandoc error: {e}")
    finally:
        if os.path.exists(tmp_tex): os.remove(tmp_tex)
        if os.path.exists(tmp_docx): os.remove(tmp_docx)
